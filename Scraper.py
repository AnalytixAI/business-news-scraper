import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

# URL of the business news website
URL = "https://www.moneycontrol.com/news/business/"

# Add headers to mimic a real browser
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36"
}

# Send a request to the website
response = requests.get(URL, headers=headers)
if response.status_code == 200:
    print("Website fetched successfully!")
else:
    print(f"Failed części: Failed to fetch website! Status Code: {response.status_code}")
    exit()

# Parse HTML content
soup = BeautifulSoup(response.text, "html.parser")

# Find all news headline elements (Update the class name if needed)
headlines = soup.find_all("h2")

# Create a Word document
doc = Document()
doc.add_paragraph("Business News Headlines")

def add_hyperlink(paragraph, text, url):
    part = OxmlElement("w:hlink")
    part.set(qn("r:id"), url)
    run = paragraph.add_run(text)
    run.underline = True  # Making it look clickable
    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
    paragraph._element.append(part)
    return run

for idx, headline in enumerate(headlines, start=1):
    link_tag = headline.find("a")
    if link_tag and "href" in link_tag.attrs:
        title = link_tag.text.strip()
        link = link_tag["href"]
        
        # Add headline as normal text
        para = doc.add_paragraph()
        para.add_run(f"{idx}. {title}")
        
        # Add clickable link on new line with space after
        link_para = doc.add_paragraph()
        link_para.add_run("    ")  # Add some indentation
        add_hyperlink(link_para, link, link)
        link_para.add_run(" ")  # Add space after hyperlink

# Check if data is extracted
if not headlines:
    print("No news headlines found. The website structure may have changed.")
    exit()

# Save to Word file
doc.save("Business_News.docx")
print("Data saved to Business_News.docx!")